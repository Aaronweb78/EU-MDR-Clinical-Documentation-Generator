"""
DOCX text extraction using python-docx
"""
from docx import Document
from pathlib import Path
from typing import Optional, Dict, List
import logging

logger = logging.getLogger(__name__)


class DOCXExtractor:
    """Extract text from DOCX files"""

    @staticmethod
    def extract_text(file_path: Path) -> Optional[str]:
        """
        Extract all text from DOCX file

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text or None if extraction fails
        """
        try:
            doc = Document(str(file_path))
            text_content = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))

            full_text = "\n".join(text_content)
            return full_text if full_text.strip() else None

        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return None

    @staticmethod
    def extract_paragraphs(file_path: Path) -> List[str]:
        """
        Extract paragraphs as a list

        Args:
            file_path: Path to DOCX file

        Returns:
            List of paragraph texts
        """
        try:
            doc = Document(str(file_path))
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            return paragraphs

        except Exception as e:
            logger.error(f"Error extracting paragraphs from DOCX {file_path}: {e}")
            return []

    @staticmethod
    def extract_tables(file_path: Path) -> List[List[List[str]]]:
        """
        Extract tables from DOCX

        Args:
            file_path: Path to DOCX file

        Returns:
            List of tables, where each table is a list of rows
        """
        try:
            doc = Document(str(file_path))
            tables = []

            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)

            return tables

        except Exception as e:
            logger.error(f"Error extracting tables from DOCX {file_path}: {e}")
            return []

    @staticmethod
    def extract_metadata(file_path: Path) -> Dict:
        """
        Extract DOCX metadata

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary of metadata
        """
        try:
            doc = Document(str(file_path))
            core_props = doc.core_properties

            return {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'created': core_props.created,
                'modified': core_props.modified,
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }

        except Exception as e:
            logger.error(f"Error extracting DOCX metadata {file_path}: {e}")
            return {}

    @staticmethod
    def extract_structured(file_path: Path) -> Dict:
        """
        Extract text with structure preserved

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with paragraphs and tables
        """
        try:
            doc = Document(str(file_path))

            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append({
                        'text': text,
                        'style': para.style.name if para.style else 'Normal'
                    })

            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    'table_index': table_idx,
                    'data': table_data
                })

            return {
                'paragraphs': paragraphs,
                'tables': tables
            }

        except Exception as e:
            logger.error(f"Error extracting structured DOCX {file_path}: {e}")
            return {'paragraphs': [], 'tables': []}

    @staticmethod
    def get_word_count(file_path: Path) -> int:
        """
        Get approximate word count

        Args:
            file_path: Path to DOCX file

        Returns:
            Number of words
        """
        try:
            text = DOCXExtractor.extract_text(file_path)
            if text:
                return len(text.split())
            return 0
        except Exception as e:
            logger.error(f"Error getting word count {file_path}: {e}")
            return 0

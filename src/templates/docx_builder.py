"""
DOCX document builder for report export
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from typing import List, Dict
import logging

logger = logging.getLogger(__name__)


class DOCXBuilder:
    """Build formatted DOCX documents from report sections"""

    def __init__(self):
        """Initialize DOCX builder"""
        self.doc = None

    def create_document(self, title: str, device_name: str = None) -> Document:
        """
        Create a new document with title page

        Args:
            title: Document title
            device_name: Device name (optional)

        Returns:
            Document object
        """
        self.doc = Document()

        # Set default styles
        self._set_default_styles()

        # Add title page
        self._add_title_page(title, device_name)

        return self.doc

    def add_section(self, section_number: int, section_title: str, content: str):
        """
        Add a section to the document

        Args:
            section_number: Section number
            section_title: Section title
            content: Section content text
        """
        if not self.doc:
            raise RuntimeError("Document not created. Call create_document() first.")

        # Add page break before section (except first)
        if section_number > 1:
            self.doc.add_page_break()

        # Add section heading
        heading = self.doc.add_heading(f"{section_number}. {section_title}", level=1)
        heading.style.font.color.rgb = RGBColor(0, 132, 170)  # Teal color

        # Add content paragraphs
        for paragraph_text in content.split('\n\n'):
            if paragraph_text.strip():
                p = self.doc.add_paragraph(paragraph_text.strip())
                p.style = 'Normal'

    def add_table_of_contents(self):
        """Add table of contents placeholder"""
        if not self.doc:
            return

        self.doc.add_page_break()
        heading = self.doc.add_heading('Table of Contents', level=1)
        self.doc.add_paragraph('[Table of Contents will be generated when opened in Word]')
        self.doc.add_paragraph('Right-click and select "Update Field" to generate.')

    def save(self, filepath: Path):
        """
        Save document to file

        Args:
            filepath: Path to save document
        """
        if not self.doc:
            raise RuntimeError("Document not created")

        try:
            filepath.parent.mkdir(parents=True, exist_ok=True)
            self.doc.save(str(filepath))
            logger.info(f"Document saved to {filepath}")
        except Exception as e:
            logger.error(f"Error saving document: {e}")
            raise

    def _set_default_styles(self):
        """Set default document styles"""
        if not self.doc:
            return

        # Normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Heading 1
        style = self.doc.styles['Heading 1']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(16)
        font.bold = True
        font.color.rgb = RGBColor(0, 132, 170)

        # Heading 2
        style = self.doc.styles['Heading 2']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(14)
        font.bold = True

    def _add_title_page(self, title: str, device_name: str = None):
        """Add title page to document"""
        if not self.doc:
            return

        # Add spacing
        for _ in range(5):
            self.doc.add_paragraph()

        # Add title
        title_para = self.doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 132, 170)

        # Add device name if provided
        if device_name:
            self.doc.add_paragraph()
            device_para = self.doc.add_paragraph(device_name)
            device_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            device_run = device_para.runs[0]
            device_run.font.size = Pt(16)
            device_run.font.color.rgb = RGBColor(100, 100, 100)

        # Add spacing
        for _ in range(3):
            self.doc.add_paragraph()

        # Add generation info
        from datetime import datetime
        date_para = self.doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.runs[0]
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)

        # Add footer note
        footer_para = self.doc.add_paragraph()
        footer_para.add_run("\n\nGenerated with EU MDR Clinical Documentation Generator")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.runs[0].font.size = Pt(9)
        footer_para.runs[0].font.italic = True
        footer_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    def build_from_sections(
        self,
        report_title: str,
        device_name: str,
        sections: List[Dict],
        output_path: Path
    ):
        """
        Build complete document from sections

        Args:
            report_title: Report title
            device_name: Device name
            sections: List of section dictionaries with 'section_number', 'section_title', 'content'
            output_path: Path to save document
        """
        # Create document
        self.create_document(report_title, device_name)

        # Add table of contents
        self.add_table_of_contents()

        # Add all sections
        for section in sections:
            if section.get('content'):
                self.add_section(
                    section_number=section['section_number'],
                    section_title=section['section_title'],
                    content=section['content']
                )

        # Save
        self.save(output_path)

        return output_path
